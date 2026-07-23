from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(__file__).resolve().parent.parent
out_dir = root / "Docs" / "documentacion final"
out_dir.mkdir(parents=True, exist_ok=True)


def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_title(doc, title, subtitle, version):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = 24

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = 12

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(version)
    run.italic = True
    run.font.size = 11
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    return table


def add_section(doc, title, paragraphs=None, bullets=None, level=1):
    doc.add_heading(title, level=level)
    if paragraphs:
        for paragraph in paragraphs:
            add_paragraph(doc, paragraph)
    if bullets:
        add_bullets(doc, bullets)


def add_page_section(doc, title, description, bullets=None):
    doc.add_heading(title, level=2)
    add_paragraph(doc, description)
    if bullets:
        add_bullets(doc, bullets)


def get_role_screens(profile_name):
    common_pages = [
        {
            "title": "Módulo de autenticación",
            "description": "El módulo de autenticación es la puerta de entrada al sistema YESA. Controla el acceso a todas las funciones protegidas y gestiona los estados de sesión que permiten usar el carrito, favoritos, historial de pedidos y áreas administrativas.",
            "bullets": [
                "Incluye las pantallas de Inicio de Sesión y Registro que el usuario utiliza para acceder al sistema.",
                "Cuando el usuario inicia sesión correctamente, el sistema decide si debe enviarlo al catálogo o al dashboard de administración, según su rol.",
                "Si se ingresan credenciales inválidas, el sistema muestra un mensaje de error con el motivo y el usuario debe corregir los datos.",
                "Si el usuario tiene un carrito local sin iniciar sesión, este se sincroniza automáticamente al iniciar sesión para no perder los artículos seleccionados.",
                "Al ser redirigido por una acción protegida, como intentar guardar un diseño o cotizar sin sesión, el sistema vuelve a la pantalla original tras el login.",
                "El botón 'Cerrar Sesión' elimina la sesión activa y retorna al usuario a la pantalla de login o a la página pública.",
                "La barra de navegación superior cambia sus opciones según el estado de sesión: botones de acceso, perfil, pedidos o administración.",
            ],
        },
        {
            "title": "Pantalla de inicio de sesión",
            "description": "La pantalla de inicio de sesión permite al usuario autenticarse con su correo y contraseña. Es el primer paso para acceder a funciones como carrito, favoritos, cotizaciones y pedidos.",
            "bullets": [
                "Campo 'Email' para escribir la dirección de correo electrónico registrada en YESA.",
                "Campo 'Contraseña' para ingresar la clave privada del usuario.",
                "Botón 'Iniciar Sesión' para enviar la información al servidor y validar el acceso.",
                "Mientras el servidor responde, el botón muestra 'Iniciando sesión...' e indica que el proceso está en curso.",
                "Si la autenticación falla, se muestra una alerta en la pantalla con el mensaje de error exacto.",
                "Si el usuario intenta acceder a una acción protegida sin sesión, después del login se redirige a la pantalla solicitada.",
                "El enlace 'Crear cuenta nueva' lleva directamente a la pantalla de Registro, manteniendo la ruta de origen para regresar después.",
                "Si el usuario ya tenía un carrito en el navegador, el sistema informa que ese carrito se sincronizará con su cuenta al iniciar sesión.",
            ],
        },
        {
            "title": "Pantalla de registro",
            "description": "La pantalla de registro crea una cuenta nueva y permite al usuario comenzar a comprar. Incluye campos para datos personales, correo, contraseña y datos de contacto.",
            "bullets": [
                "Campo 'Nombre' para ingresar el nombre de pila del usuario.",
                "Campo 'Apellido' para ingresar el apellido o segundo nombre.",
                "Campo 'Email' para ingresar la dirección de correo que se usará como usuario.",
                "Campo 'Contraseña' para establecer una clave segura de acceso.",
                "Campo 'Confirmar Contraseña' para repetir la contraseña y garantizar que no haya errores tipográficos.",
                "Campo 'Teléfono' opcional para agregar un número de contacto con formato de 10 dígitos comenzando con 3.",
                "Campo 'Dirección' opcional para información de envío o contacto adicional.",
                "Botón 'Crear Cuenta' envía el formulario y, si los datos son correctos, crea el usuario en el sistema.",
                "Si la contraseña es muy corta, no coincide o el email es inválido, el formulario muestra un mensaje de error y no avanza.",
                "Si el usuario ya existe o hay un problema en el registro, el sistema muestra la razón del error para que el usuario pueda corregirla.",
                "El enlace 'Iniciar Sesión' permite regresar al login si el usuario ya posee cuenta.",
                "Cuando el registro es exitoso, el usuario se redirige automáticamente al catálogo para continuar comprando.",
            ],
        },
        {
            "title": "Menú principal y navegación",
            "description": "La barra superior es el punto central de navegación y muestra enlaces diferentes según el usuario. Permite ir al inicio, al catálogo, al carrito, a las compras, al perfil y a las secciones administrativas.",
            "bullets": [
                "Botón 'Inicio' devuelve al usuario a la página principal donde puede ver información general del sistema.",
                "Botón 'Catálogo' lleva a la lista completa de productos disponibles para revisar y comprar.",
                "Botón 'Carrito' muestra el resumen de productos seleccionados y permite avanzar al pago.",
                "Botón 'FAQ' abre un modal con preguntas frecuentes y acceso a soporte para dudas sobre el uso de la plataforma.",
                "Los usuarios no autenticados ven los botones 'Iniciar Sesión' y 'Registro' en la barra superior.",
                "Los usuarios autenticados ven su nombre, el enlace a 'Mi Perfil' y el botón 'Cerrar Sesión'.",
                "Los clientes autenticados ven además el botón 'Mis Pedidos' para consultar el historial de compras.",
                "Los clientes autenticados también pueden acceder a 'Mis Consultas' desde el menú para revisar su comunicación con soporte.",
                "Los auxiliares y administradores ven el menú 'Administración' con accesos rápidos a Dashboard, Categorías, Subcategorías, Productos y Pedidos.",
                "Los administradores ven en el menú de administración accesos adicionales a Usuarios y Cotizaciones.",
                "En la ruta de catálogo, aparece la barra de búsqueda para filtrar productos y un menú desplegable de categorías para refinar resultados.",
                "Este menú permite identificar rápidamente qué funciones tienes disponible según tu rol y evita acceder a secciones no autorizadas.",
            ],
        },
        {
            "title": "Catálogo",
            "description": "El catálogo es la página principal de exploración de productos. Presenta todas las opciones disponibles en tarjetas, permite buscar por texto y filtrar por categoría, además de llevar directamente al detalle de cada producto.",
            "bullets": [
                "Cada tarjeta de producto muestra la imagen, el nombre, el precio y la disponibilidad en stock.",
                "El botón 'Ver' abre la ficha completa del producto con descripción, imágenes y opciones de compra.",
                "El botón 'Agregar al carrito' añade el producto al carrito sin salir del catálogo.",
                "El botón 'Personalizar' lleva al usuario al módulo de personalización para crear un diseño propio basado en el producto.",
                "La barra de búsqueda filtra productos mientras se escribe y actualiza la URL para conservar el filtro.",
                "El filtro de categorías permite ver solo los productos que pertenecen a la categoría seleccionada.",
                "Si no se encuentran resultados, la pantalla muestra un mensaje indicando que no hay productos y sugiere modificar la búsqueda o filtros.",
                "El catálogo también puede mostrar datos de stock o precios especiales según la configuración del producto.",
            ],
        },
        {
            "title": "Detalle de producto",
            "description": "La ficha de producto muestra información completa: galería de imágenes, descripción, precio, stock y acciones para comprar o personalizar. Es el punto donde el cliente decide si agrega el producto al carrito.",
            "bullets": [
                "Las miniaturas permiten cambiar la imagen principal para ver el producto desde ángulos diferentes.",
                "El selector de cantidad permite indicar cuántas unidades se desean comprar antes de agregar al carrito.",
                "El botón 'Agregar al carrito' añade la cantidad seleccionada al carrito y actualiza el total.",
                "El botón de corazón permite marcar el producto como favorito o quitarlo de favoritos si ya estaba marcado.",
                "El botón 'Personalizar' abre el editor 3D para diseñar el producto con colores, texturas y texto.",
                "Si el usuario intenta usar favoritos o cotizaciones sin iniciar sesión, el sistema solicita el login y regresa luego a la ficha.",
                "La ficha también muestra el estado del stock y ayuda a evitar comprar productos agotados.",
            ],
        },
        {
            "title": "Carrito",
            "description": "La página de carrito muestra todos los productos que el usuario seleccionó para comprar, permite modificar cantidades, eliminar artículos y avanzar hacia el pago.",
            "bullets": [
                "Cada producto en el carrito muestra nombre, precio unitario, cantidad y subtotal.",
                "El botón '-' reduce la cantidad de ese producto en una unidad.",
                "El botón '+' aumenta la cantidad de ese producto en una unidad.",
                "El botón 'Eliminar' quita el producto seleccionado del carrito.",
                "El botón 'Vaciar carrito' elimina todos los productos a la vez y deja el carrito vacío.",
                "Si el carrito está vacío, el sistema muestra un mensaje invitando al usuario a volver al catálogo.",
                "El botón 'Proceder al Pago' lleva al usuario a la pantalla de checkout para finalizar la compra.",
                "Si el usuario no está autenticado, en lugar de pagar se muestra 'Iniciar Sesión para Pagar'.",
                "El carrito también muestra totales y resumen de costos antes de pagar.",
            ],
        },
        {
            "title": "Checkout",
            "description": "La pantalla de checkout recopila los datos de envío, el método de pago y las notas adicionales. Es el último paso para confirmar el pedido y enviarlo al sistema.",
            "bullets": [
                "El campo 'Dirección de Envío' solicita la dirección completa donde se entregará el pedido.",
                "El campo 'Teléfono de Contacto' debe contener un número válido para que el equipo pueda comunicarse con el cliente.",
                "El selector 'Método de Pago' permite elegir entre efectivo, tarjeta de crédito/débito o transferencia bancaria.",
                "El campo 'Notas Adicionales' es opcional y sirve para dejar instrucciones específicas de entrega.",
                "El botón 'Confirmar Pedido' envía la orden al sistema y crea el pedido asociado al usuario.",
                "El botón 'Volver al Carrito' permite regresar y revisar los productos antes de confirmar.",
                "Si el pedido proviene de una cotización, el botón cambia a 'Volver a Cotizaciones' para regresar al detalle de la cotización.",
                "Después de confirmar, se espera una pantalla de pedido confirmado con resumen y seguimiento.",
            ],
        },
        {
            "title": "Personalización",
            "description": "El módulo de personalización ofrece herramientas visuales para diseñar un producto en modo 3D, incluyendo colores, texturas, texto y controles de vista. Permite guardar diseños, cotizarlos y compartirlos con el equipo.",
            "bullets": [
                "Botón '-' reduce el nivel de zoom del modelo para ver el diseño completo.",
                "Botón '+' aumenta el nivel de zoom para centrar detalles específicos.",
                "El botón de rotación cambia si el modelo gira automáticamente o se detiene.",
                "El botón de pantalla completa expande el visor para una vista más cómoda del diseño.",
                "Los selectores de color aplican tonos a Interior, Base, Exterior y Asa del modelo.",
                "El botón 'Agregar texto' abre el editor de texto que permite seleccionar tipografía, tamaño y color.",
                "El botón 'Elegir archivo' abre el selector de archivos para subir una imagen como textura (JPG, PNG, GIF hasta 25MB).",
                "El botón 'Limpiar' borra la textura aplicada y restaura los colores básicos del modelo.",
                "En el editor de texto, el botón 'Aplicar texto' guarda los cambios en el diseño y los muestra en la vista 3D.",
                "Los controles de movimiento permiten ajustar la posición de la textura cargada en el modelo.",
                "Los controles de escala permiten aumentar o disminuir el tamaño de la textura aplicada.",
                "El campo 'Nombre del diseño para la cotización' sirve para identificar el diseño dentro de las cotizaciones.",
                "El campo 'Notas para la cotización' se usa para explicar detalles adicionales al equipo de producción.",
                "El botón 'Cotizar producto' envía el diseño al sistema para obtener un precio estimado.",
                "El botón 'Guardar diseño' guarda el diseño en la lista de diseños guardados para editarlo después.",
                "El botón 'Compartir' copia un enlace o permite compartir el diseño con otras personas.",
                "Si el usuario no está autenticado, aparece un aviso indicando que debe iniciar sesión para guardar o cotizar el diseño.",
                "El módulo conserva el diseño pendiente mientras el usuario inicia sesión, para que no se pierda el trabajo realizado.",
            ],
        },
    ]

    if profile_name == "Usuario sin registrar":
        return [
            *common_pages,
            {
                "title": "Acceso y registro",
                "description": "Los usuarios sin registrar pueden ver la plataforma pero necesitan iniciar sesión para completar acciones clave.",
                "bullets": [
                    "Formulario de inicio de sesión con correo electrónico y contraseña.",
                    "Enlace para crear una cuenta nueva si todavía no está registrado.",
                    "Intentos de guardar diseños, cotizar o pagar muestran la necesidad de autenticarse.",
                    "Se puede navegar en el catálogo, pero no se puede finalizar un pedido sin iniciar sesión.",
                ],
            },
            {
                "title": "Restricciones principales",
                "description": "Estas funciones visuales están limitadas hasta que el usuario se registre e inicie sesión.",
                "bullets": [
                    "Agregar al carrito es visible, pero el checkout requiere autenticación.",
                    "Favoritos no se cargan sin sesión y redirigen al login si el usuario no está autenticado.",
                    "Guardar diseños o cotizar desde personalización no es posible sin iniciar sesión.",
                    "Acceso a 'Mis pedidos', 'Mis cotizaciones' y 'Perfil' está restringido.",
                ],
            },
        ]

    client_pages = [
        *common_pages,
        {
            "title": "Favoritos",
            "description": "La página de favoritos muestra los productos guardados por el cliente.",
            "bullets": [
                "Botón 'Ver Detalles' para abrir la ficha del producto favorito.",
                "Botón de corazón interno en la ficha para eliminar ese producto de favoritos.",
                "Si no hay favoritos, hay un botón 'Explorar Catálogo' para regresar al catálogo.",
            ],
        },
        {
            "title": "Mis pedidos",
            "description": "El cliente puede revisar el historial de compras y el estado de cada pedido.",
            "bullets": [
                "Botón 'Seguir Comprando' para volver al catálogo desde mis pedidos.",
                "Lista de pedidos con ID, fecha, estado y total.",
                "Botón 'Ver Detalle' para ver los productos y el seguimiento de cada pedido.",
            ],
        },
        {
            "title": "Mis cotizaciones",
            "description": "El cliente accede a sus solicitudes de cotización y puede avanzar a pedido cuando estén cotizadas.",
            "bullets": [
                "Tarjetas de cotización que indican estado, subtotal estimado y fecha.",
                "Botón 'Ver detalle' para abrir la página de la cotización concreta.",
                "Botón 'Realizar pedido' habilitado solo cuando el estado es 'cotizado'.",
                "Si la cotización no está cotizada, al presionar el botón se muestra un mensaje de espera.",
                "Cada cotización muestra el usuario, el número de diseños y el estado actual para facilitar su seguimiento.",
            ],
        },
        {
            "title": "Mis Consultas",
            "description": "Esta sección permite al cliente ver el seguimiento de las consultas enviadas al equipo de soporte.",
            "bullets": [
                "Lista de consultas con asunto, fecha y estado (pendiente, revisado, respondido o cerrado).",
                "Cada tarjeta muestra el mensaje original y la respuesta cuando está disponible.",
                "Si la consulta se marca como respondida, se muestra un panel con la respuesta del equipo.",
                "La consulta cerrada indica que el caso ya fue atendido y no requiere más seguimiento.",
                "El acceso a esta sección ayuda a llevar un registro ordenado de las solicitudes de soporte.",
            ],
        },
        {
            "title": "Diseños guardados",
            "description": "La sección de diseños guardados muestra los diseños 3D creados y permite cotizarlos o editarlos.",
            "bullets": [
                "Botón 'Crear nuevo diseño' para regresar al módulo de personalización.",
                "Botón 'Eliminar todo' para eliminar todos los diseños guardados localmente.",
                "Checkbox de selección para elegir diseños y cotizarlos en lote.",
                "Botón 'Cotizar seleccionados' para enviar los diseños marcados a cotización.",
                "Botón 'Editar diseño' para recuperar un diseño guardado y seguir personalizándolo.",
                "Botón 'Eliminar diseño' para sacar un diseño individual de la lista.",
                "Si no hay diseños guardados, la página muestra un mensaje y un botón 'Crear nuevo diseño'.",
                "Si el usuario no está autenticado, al intentar cotizar aparece un aviso y se redirige al login.",
            ],
        },
        {
            "title": "Perfil",
            "description": "El perfil permite actualizar datos personales y cambiar la contraseña de la cuenta.",
            "bullets": [
                "Campo 'Nombre' para modificar el nombre del usuario.",
                "Campo 'Email' para actualizar el correo electrónico.",
                "Campo 'Teléfono' para registrar el número de contacto.",
                "Campo 'Dirección' para la dirección de envío o contacto.",
                "Botón 'Guardar cambios' para actualizar los datos del perfil.",
                "Sección 'Cambiar contraseña' con campos: contraseña actual, nueva contraseña y confirmar nueva contraseña.",
                "Botón 'Cambiar contraseña' para enviar la solicitud de actualización de clave.",
                "El sistema valida la información antes de guardar y muestra mensajes de éxito o error según la respuesta.",
            ],
        },
    ]

    if profile_name == "Cliente":
        return client_pages

    auxiliar_pages = [
        *client_pages,
        {
            "title": "Dashboard administrativo",
            "description": "El dashboard ofrece una vista general del estado de la tienda y permite abrir rápidamente las funciones administrativas más importantes.",
            "bullets": [
                "Botón 'Actualizar datos' para refrescar métricas y números de estado.",
                "Tarjetas principales de KPI para Productos, Pedidos, Usuarios, Ventas Totales, Categorías, Subcategorías, Cotizaciones y Soporte.",
                "Clic en una tarjeta de KPI para ir directamente a la sección relacionada.",
                "Sección 'Distribución de Pedidos' con resúmenes por estado y porcentaje del total.",
                "Sección 'Actividad Reciente' con los últimos pedidos del sistema y acceso directo a 'Ver detalle'.",
                "Sección 'Acciones Rápidas' con accesos a 'Nuevo Producto', 'Nueva Categoría', 'Subcategorías', 'Gestionar Pedidos', 'Nuevo Usuario' y 'Cotizaciones'.",
            ],
        },
        {
            "title": "Gestión de pedidos (auxiliar)",
            "description": "Permite revisar los pedidos, aplicar filtros y cambiar su estado desde una vista centralizada.",
            "bullets": [
                "Botón 'Volver' para regresar al dashboard administrativo.",
                "Botón 'Exportar' con menú desplegable para descargar en PDF o Excel los pedidos filtrados.",
                "Campo 'Buscar cliente' para localizar pedidos por nombre o correo del cliente.",
                "Filtros por 'Estado', 'Fecha inicio', 'Fecha fin' y opción para ordenar por fecha.",
                "Botón 'Ver detalle' para abrir el modal con información completa del pedido y sus productos.",
                "Botones en cada fila para cambiar el estado: 'En Proceso', 'Cancelar', 'Enviar' o 'Entregar', según el estado actual.",
                "En el modal de detalle se muestran también los botones de cambio de estado: 'Pendiente', 'En Proceso', 'Enviado', 'Entregado' y 'Cancelado'.",
            ],
        },
        {
            "title": "Gestión de categorías",
            "description": "La sección de categorías permite crear, editar, activar, desactivar y eliminar categorías del catálogo.",
            "bullets": [
                "Botón 'Exportar' para generar un reporte en PDF o Excel de las categorías filtradas.",
                "Menú de acción masiva con 'Activar todas' y 'Desactivar todas' para las categorías visibles con filtros.",
                "Botón 'Nueva Categoría' para abrir el modal de creación.",
                "Botones 'Editar' en cada fila para modificar nombre, descripción y estado.",
                "Botón 'Eliminar' para borrar una categoría con confirmación.",
                "Botón 'Activar' o 'Desactivar' para cambiar el estado de cada categoría individualmente.",
                "Filtros de búsqueda, estado y orden por nombre, además del botón 'Limpiar filtros'.",
            ],
        },
        {
            "title": "Gestión de subcategorías",
            "description": "La gestión de subcategorías controla la creación y organización de las subcategorías del catálogo.",
            "bullets": [
                "Botón 'Exportar' para guardar la lista en PDF o Excel.",
                "Menú de acción masiva con 'Activar todas' y 'Desactivar todas' para las subcategorías filtradas.",
                "Botón 'Nueva Subcategoría' para abrir el formulario de registro.",
                "Botón 'Editar' para cambiar nombre, descripción, categoría asociada y estado.",
                "Botón 'Eliminar' para borrar la subcategoría con confirmación.",
                "Botón 'Activar' o 'Desactivar' para cambiar el estado individualmente.",
                "Filtros por búsqueda, categoría, estado y orden por nombre, con botón 'Limpiar'.",
            ],
        },
        {
            "title": "Gestión de productos",
            "description": "La sección de productos permite crear, editar y controlar la disponibilidad de los artículos del catálogo.",
            "bullets": [
                "Botón 'Exportar' para descargar un reporte de productos filtrados en PDF o Excel.",
                "Botón 'Nuevo Producto' para abrir el modal de creación.",
                "Filtros por búsqueda, categoría, subcategoría, rango de precio y orden por nombre.",
                "Botón 'Editar' para modificar nombre, descripción, precio, stock, categoría, subcategoría, activo e imagen.",
                "Botón 'Eliminar' para quitar un producto del catálogo.",
                "Botón 'Activar' o 'Desactivar' para cambiar la visibilidad del producto.",
                "Botón de acción masiva para activar o desactivar todos los productos filtrados.",
                "Sección de 'Vista Previa' para adjuntar una imagen del producto desde el modal de edición o creación.",
            ],
        },
        {
            "title": "Soporte",
            "description": "La sección de soporte permite revisar mensajes recibidos y responderlos desde la plataforma.",
            "bullets": [
                "Campos de filtro por búsqueda, estado y asunto para ubicar tickets específicos.",
                "Botón 'Ver detalle' para abrir el contenido completo del mensaje de contacto.",
                "Botón 'Responder' para escribir y enviar una respuesta al usuario.",
                "Botón 'Cerrar ticket' para finalizar el caso cuando se resuelve.",
                "Cada mensaje puede tener estado 'Pendiente', 'Revisado', 'Respondido' o 'Cerrado'.",
            ],
        },
    ]

    if profile_name == "Auxiliar":
        return auxiliar_pages

    administrador_pages = [
        *auxiliar_pages,
        {
            "title": "Dashboard administrativo",
            "description": "El dashboard ofrece una vista general de métricas clave y accesos rápidos a las secciones administrativas.",
            "bullets": [
                "Botón 'Actualizar datos' para refrescar métricas y números de estado.",
                "Tarjetas de KPI con Totales de Categorías, Subcategorías, Productos, Usuarios y Pedidos.",
                "Tarjetas de KPI con acceso directo a Categorías, Subcategorías, Cotizaciones y Soporte.",
                "Clic en una tarjeta de KPI de producto, pedido o usuario para ir directamente a esa sección.",
                "Se muestran métricas de ventas totales, stock bajo y mensajes pendientes de soporte.",
                "Botones rápidos 'Ver productos' y 'Ver usuarios' para ir directamente a esas secciones.",
                "Acciones rápidas disponibles para Nuevo Producto, Nueva Categoría, Subcategorías, Gestionar Pedidos y Cotizaciones.",
            ],
        },
        {
            "title": "Gestión de usuarios",
            "description": "El administrador administra cuentas, roles y accesos de los usuarios de YESA.",
            "bullets": [
                "Botón 'Nuevo Usuario' para crear un registro con nombre, apellido, email, contraseña, teléfono, dirección, rol y estado.",
                "Botón 'Editar' para actualizar los datos del usuario y dejar la contraseña vacía si no se desea cambiar.",
                "Botón 'Activar' o 'Desactivar' para cambiar el estado del usuario sin eliminarlo.",
                "Botón 'Eliminar' para borrar el usuario del sistema con confirmación previa.",
                "Botones de exportación 'Exportar' con opciones PDF y Excel para descargar los usuarios filtrados.",
                "Filtros por 'Buscar por nombre o email', 'Rol' y 'Estado', y botón 'Limpiar filtros' para reiniciar la búsqueda.",
                "Paginación con botones 'Anterior', 'Siguiente' y página actual 'Página x de y' para navegar por los resultados.",
                "Mensaje de advertencia si intenta desactivar su propia cuenta: 'No puedes desactivarte a ti mismo'.",
            ],
        },
        {
            "title": "Gestión de cotizaciones",
            "description": "El administrador revisa presupuestos, asigna precios y actualiza estados de cotizaciones generadas desde personalización.",
            "bullets": [
                "Cada tarjeta muestra ID, nombre de cotización, usuario, estado y precio estimado.",
                "Campo de entrada para 'Asignar precio' a la cotización y botón 'Asignar' para guardar el valor.",
                "Si el precio está pendiente, la tarjeta muestra 'Pendiente' hasta que se asigne un valor.",
                "Botón 'Ver detalle' para revisar la cotización completa y ver la información del diseño.",
                "Botón 'Volver al dashboard' para regresar fácilmente a la vista principal de administración.",
                "Solo usuarios administradores pueden acceder a la sección de cotizaciones.",
            ],
        },
        {
            "title": "Detalle de cotización (admin)",
            "description": "El administrador revisa un diseño cotizado y su información antes de validar el presupuesto.",
            "bullets": [
                "Botón 'Volver' para regresar a la lista de cotizaciones.",
                "Etiqueta de estado con colores para 'pendiente', 'cotizado', 'aceptado' o 'rechazado'.",
                "Vista previa 3D del diseño y panel con datos de colores, textura, texto y modelo usado.",
                "Resumen con precio, usuario, fecha, notas y estado de la cotización.",
                "Botón 'Volver al dashboard' disponible también desde esta vista para regresar al panel principal.",
            ],
        },
    ]

    if profile_name == "Administrador":
        return administrador_pages

    return common_pages


def build_profile_manual(profile_name, role_label, role_description, focus_points, faq_items):
    doc = Document()
    add_title(
        doc,
        f"Manual de Usuario YESA - {profile_name}",
        f"Guía detallada para {role_label.lower()} del sistema YESA",
        "Versión 2.0 • 2026 • Elaborado para usuarios sin conocimientos técnicos",
    )

    add_paragraph(doc, "ÍNDICE")
    add_bullets(doc, [
        "1. Descripción del sistema",
        "2. Objetivo y alcance",
        "3. Funcionalidades principales",
        "4. Requisitos básicos",
        "5. Registro e inicio de sesión",
        "6. Navegación por la interfaz",
        "7. Pantallas clave y botones",
        "8. Gestión del catálogo y productos",
        "9. Gestión del carrito, favoritos y pedidos",
        "10. Personalización y cotizaciones",
        "11. Restricciones y responsabilidades del rol",
        "12. Soporte, seguridad y resolución de problemas",
        "13. Preguntas frecuentes",
        "14. Glosario",
        "15. Anexos",
    ])

    add_section(doc, "1. Descripción del sistema", [
        "YESA es una plataforma de comercio digital enfocada en la venta de productos, la gestión de pedidos y la personalización de diseños.",
        "El sistema está pensado para que cualquier persona, incluso sin conocimientos técnicos, pueda usarlo con claridad y seguridad.",
        f"Este documento está orientado especialmente a {role_label.lower()} del sistema YESA. {role_description}",
    ])

    add_section(doc, "1.1 Objetivo", [
        "El objetivo de este manual es enseñar de forma clara cómo usar YESA, comprender sus funciones principales y evitar errores durante el proceso diario.",
        "También busca explicar por qué algunas acciones están restringidas, para que el usuario comprenda la importancia de los permisos y la seguridad del sistema.",
    ], level=2)

    add_section(doc, "1.2 Alcance", [
        "Este manual cubre desde la creación de cuenta hasta la realización de compras, la personalización de productos, la revisión de pedidos y la solicitud de ayuda por soporte.",
        "Se incluyen además aspectos de seguridad, buenas prácticas, uso de roles y situaciones frecuentes que pueden presentarse durante la interacción con la plataforma.",
    ], level=2)

    add_section(doc, "2. Funcionalidades principales", bullets=[
        "Explorar productos y categorías.",
        "Usar la búsqueda para encontrar artículos particulares.",
        "Agregar productos al carrito y revisar el pedido antes de pagar.",
        "Guardar artículos como favoritos para revisarlos luego.",
        "Crear diseños personalizados y generar cotizaciones.",
        "Consultar el estado de pedidos, compras y solicitudes de soporte.",
        "Administrar o apoyar tareas según el rol asignado.",
        "Recibir ayuda a través del módulo de soporte cuando se presente algún problema.",
    ])

    add_section(doc, "2.1 Mapa del sistema", bullets=[
        "Módulo de autenticación: pantalla de Inicio de Sesión con Email y Contraseña, botón 'Iniciar Sesión', mensajes de error y enlace al registro; pantalla de Registro con campos Nombre, Apellido, Email, Contraseña, Confirmar Contraseña, Teléfono, Dirección y botón 'Crear Cuenta'.",
        "Módulo de catálogo: página de listado de productos con búsqueda, filtros por categorías, tarjetas de producto, botones 'Ver', 'Agregar al carrito' y 'Personalizar'.",
        "Módulo de carrito: página con lista de artículos, botones '-', '+', 'Eliminar', 'Vaciar carrito' y 'Proceder al pago'.",
        "Módulo de checkout: formulario de pago con campos de dirección, teléfono, método de pago, notas adicionales, botón 'Confirmar Pedido' y opción para volver al carrito o a cotizaciones.",
        "Módulo de pedidos: historial de pedidos con estados, botones 'Ver Detalle' y 'Seguir Comprando' para revisar y continuar comprando.",
        "Módulo de personalización: editor de diseños 3D con zoom, rotación, pantalla completa, selección de colores, carga de texturas, texto overlay, cotización, guardado y compartición.",
        "Módulo de soporte: consultas y tickets con formulario de contacto, mensajes enviados y respuestas desde el panel de soporte.",
        "Módulo administrativo: menú de administración con accesos a Dashboard, Categorías, Subcategorías, Productos, Pedidos, Usuarios y Cotizaciones según el rol.",
    ], level=2)

    add_section(doc, "3. Requisitos básicos", bullets=[
        "Debe tener un navegador actualizado como Chrome, Edge o Firefox.",
        "Debe disponer de conexión a internet estable.",
        "Debe contar con credenciales válidas si ya tiene una cuenta.",
        "Si la plataforma se usa en entorno local, deben estar activos el backend y la base de datos.",
        "Debe revisar que su rol le permita ejecutar las acciones que desea realizar.",
    ])

    add_section(doc, "4. Registro e inicio de sesión", bullets=[
        "En la pantalla de Inicio de Sesión, ingrese su Email y Contraseña en los campos correspondientes.",
        "Presione el botón 'Iniciar Sesión' para validar sus credenciales y acceder al sistema.",
        "Si la cuenta no existe o la contraseña es incorrecta, el sistema muestra un mensaje de error en rojo.",
        "Si no tiene cuenta, presione el botón 'Crear cuenta nueva' para ir a la pantalla de Registro.",
        "En el Registro, complete Nombre, Apellido, Email, Contraseña, Confirmar Contraseña, Teléfono y Dirección.",
        "La contraseña debe tener al menos 6 caracteres y los campos de contraseña deben coincidir.",
        "Use el botón 'Crear Cuenta' para registrar su usuario; el sistema lo enviará al catálogo cuando el registro sea exitoso.",
        "Desde la pantalla de Registro, puede volver al Login presionando el botón 'Iniciar Sesión'.",
        "Si hay un carrito local, el sistema sincroniza los artículos guardados cuando inicia sesión o se registra.",
        "Si no puede iniciar sesión, revise su correo, contraseña y conexión, y use soporte si el problema persiste.",
    ])

    add_section(doc, "5. Navegación por la interfaz", bullets=[
        "La barra de navegación superior incluye: Inicio, Catálogo, Carrito y FAQ.",
        "Para usuarios no autenticados, la barra muestra los botones 'Iniciar Sesión' y 'Registro'.",
        "Para usuarios autenticados, la barra muestra el nombre del usuario, 'Mi Perfil' y 'Cerrar Sesión'.",
        "Los clientes ven adicionalmente el botón 'Mis Pedidos'.",
        "Los auxiliares y administradores ven el menú 'Administración' con accesos a Dashboard, Categorías, Subcategorías, Productos y Pedidos.",
        "Los administradores ven además accesos a Usuarios y Cotizaciones dentro del menú de administración.",
        "En el catálogo, la barra de búsqueda filtra productos y actualiza la URL para mostrar resultados.",
        "El filtro de categorías en el menú superior permite seleccionar una categoría y ver solo sus productos.",
    ])

    doc.add_heading("6. Pantallas clave y botones", level=1)
    for page in get_role_screens(profile_name):
        add_page_section(doc, page["title"], page["description"], page.get("bullets"))

    add_section(doc, "7. Gestión del catálogo y productos", bullets=[
        "Use la página de catálogo para ver todas las tarjetas de producto con imagen, nombre, precio y stock.",
        "Escriba en la barra de búsqueda para filtrar productos por palabra clave; el sistema actualiza la lista y la URL.",
        "Seleccione una categoría en el menú superior para ver solo productos de esa categoría.",
        "Desde cada tarjeta, use el botón 'Ver' para abrir la ficha del producto y revisar detalles completos.",
        "Use el botón 'Agregar al carrito' en la tarjeta o en la ficha para enviar el producto al carrito.",
        "Use el botón 'Personalizar' para crear un diseño propio desde el producto seleccionado.",
        "Revise el producto en detalle antes de agregarlo al carrito para asegurar precio y disponibilidad.",
    ])

    add_section(doc, "8. Gestión del carrito, favoritos y pedidos", bullets=[
        "Seleccione un producto y agrégalo al carrito cuando lo quiera comprar.",
        "Revise los productos agregados antes de confirmar la compra.",
        "Modifique cantidades o elimine artículos si cambió de opinión.",
        "Use favoritos cuando quiera guardar un producto para revisarlo más tarde.",
        "Confirme la compra solo cuando esté seguro de los productos, cantidades y total.",
        "Revise la sección de pedidos para ver el estado del proceso de compra.",
    ])

    add_section(doc, "9. Personalización y cotizaciones", bullets=[
        "Acceda a la pantalla de personalización para modificar un producto con vista 3D.",
        "Use los botones '-' y '+' para ajustar el zoom, y el botón de rotación para iniciar o detener la rotación automática.",
        "Active la pantalla completa para ver el diseño en modo ampliado.",
        "Cambie los colores de Interior, Base, Exterior y Asa con los selectores de color.",
        "Use 'Agregar texto' para abrir el editor de texto, seleccione fuente, tamaño y color y confirme con 'Aplicar texto'.",
        "Use 'Elegir archivo' para subir una textura propia en JPG, PNG o GIF de hasta 25MB.",
        "Use los controles de posición para mover la textura cargada arriba, abajo, izquierda o derecha.",
        "Use los botones de escala para aumentar o disminuir el tamaño de la textura cargada.",
        "Use 'Limpiar' para eliminar la textura actual y volver a los colores básicos.",
        "Complete el campo 'Nombre del diseño para la cotización' y 'Notas para la cotización' antes de cotizar.",
        "Presione 'Cotizar producto' para generar una cotización basada en el diseño actual.",
        "Use 'Guardar diseño' para conservar el diseño en la lista de diseños guardados.",
        "Use 'Compartir' para copiar un enlace e invitar a otra persona a ver el diseño.",
        "Si no está autenticado, el sistema muestra un aviso y redirige al login conservando el diseño.",
    ])

    add_section(doc, "10. Restricciones y responsabilidades del rol", [role_description], bullets=focus_points)
    add_paragraph(doc, "Estas restricciones existen para proteger la información, evitar cambios accidentales y mantener orden dentro del sistema. También ayudan a reducir errores que puedan afectar ventas, inventario o la confianza del cliente.")

    add_section(doc, "10.1 Resumen de permisos por rol", level=2)
    add_table(doc, ["Rol", "Qué puede hacer", "Por qué tiene esa limitación"], [
        ("Cliente", "Comprar, revisar pedidos, usar carrito y favoritos", "Solo debe interactuar con su propia experiencia de compra"),
        ("Auxiliar", "Apoyar tareas administrativas básicas y revisar operaciones", "Tiene acceso de apoyo, pero no debe realizar cambios críticos"),
        ("Administrador", "Gestionar productos, categorías, usuarios, pedidos y soporte", "Es el rol con mayor responsabilidad y control del sistema"),
        ("Usuario sin registrar", "Navegar catálogo y conocer la plataforma", "No tiene cuenta activa ni acceso a funciones protegidas como comprar o guardar favoritos"),
    ])

    add_section(doc, "11. Soporte, seguridad y resolución de problemas", bullets=[
        "Use soporte cuando tenga dudas, no pueda iniciar sesión, detecte un error o necesite ayuda con una compra.",
        "Explique el problema con claridad, indique el producto o pedido y describa la pantalla donde ocurrió el error.",
        "Nunca comparta su contraseña con otras personas.",
        "Cierre sesión si está usando un equipo compartido.",
        "Revise que los datos personales estén correctos antes de comprar o modificar información.",
        "Si detecta actividad sospechosa, reporte el caso de inmediato al soporte o al administrador.",
    ])

    add_section(doc, "12. Preguntas frecuentes")
    for item, response in faq_items:
        doc.add_heading(item, level=2)
        add_paragraph(doc, response)

    add_section(doc, "13. Glosario", bullets=[
        "Catálogo: lista de productos visibles para el usuario.",
        "Carrito: espacio donde se colocan los productos elegidos para comprar.",
        "Pedido: compra confirmada dentro del sistema.",
        "Cotización: propuesta de diseño o precio relacionada con una personalización.",
        "Stock: cantidad disponible de un producto.",
        "Favorito: producto guardado para revisarlo más tarde.",
        "Soporte: canal de ayuda para reportar problemas o dudas.",
        "Rol: conjunto de permisos que determina qué acciones puede ejecutar el usuario.",
    ])

    add_section(doc, "14. Anexos", bullets=[
        "Guarde este manual como referencia para recordar procedimientos cuando los necesite.",
        "Revise el estado de los pedidos de forma periódica si está esperando una entrega o una aprobación.",
        "En caso de un problema grave, contacte al administrador o al equipo responsable.",
        "Siempre es preferible resolver una duda por soporte antes de intentar corregir un error sin información.",
    ])

    add_paragraph(doc, "Fin del manual. Este documento debe usarse como guía práctica, útil y sencilla durante el uso diario de YESA.")
    return doc


profiles = [
    {
        "name": "Cliente",
        "role_label": "cliente",
        "role_description": "El cliente usa YESA para buscar productos, agregarlos al carrito, realizar compras y revisar el estado de los pedidos.",
        "focus_points": [
            "Puede navegar por el catálogo, revisar productos y comprar con su cuenta.",
            "Puede guardar productos como favoritos y usar el carrito antes de confirmar la compra.",
            "No debe acceder a funciones administrativas ni modificar datos sensibles del sistema.",
        ],
        "faq_items": [
            ("¿Cómo puedo comprar un producto?", "Seleccione un producto, agrégelo al carrito, revise el pedido y confirme la compra."),
            ("¿Puedo cancelar un pedido?", "Debe revisar la condición del pedido y, si aplica, solicitar asistencia por soporte."),
            ("¿Qué pasa si no puedo iniciar sesión?", "Revise su correo, contraseña y estado de la cuenta; si el problema continúa, use soporte."),
        ],
    },
    {
        "name": "Auxiliar",
        "role_label": "auxiliar",
        "role_description": "El auxiliar apoya la operación de la tienda revisando órdenes, colaborando en gestión limitada y ayudando a mantener organizada la información.",
        "focus_points": [
            "Puede revisar información relevante del sistema y colaborar con tareas administrativas básicas.",
            "Debe evitar cambios críticos como eliminar usuarios o modificar datos sensibles sin autorización.",
            "Su función es apoyar, no tomar decisiones de alto impacto sin supervisión.",
        ],
        "faq_items": [
            ("¿Qué tareas puede hacer un auxiliar?", "Puede apoyar con revisiones, consultas, gestión de pedidos y tareas limitadas del sistema."),
            ("¿Puede modificar productos?", "Solo si tiene autorización y el procedimiento lo permite; su rol es de apoyo, no de administración completa."),
            ("¿Qué debe hacer si detecta un error importante?", "Debe reportarlo al administrador o al equipo responsable para evitar afectar la operación."),
        ],
    },
    {
        "name": "Administrador",
        "role_label": "administrador",
        "role_description": "El administrador gestiona la tienda completa, incluyendo productos, categorías, pedidos, usuarios, soporte y otros procesos clave del sistema.",
        "focus_points": [
            "Puede crear, modificar y desactivar categorías, subcategorías y productos.",
            "Puede revisar pedidos, gestionar usuarios y supervisar el estado general del sistema.",
            "Debe actuar con responsabilidad porque sus cambios pueden afectar ventas, inventario y confianza del cliente.",
        ],
        "faq_items": [
            ("¿Qué hace el administrador?", "Gestiona la operación completa de la tienda, los usuarios, los productos y los pedidos."),
            ("¿Qué debe revisar antes de publicar un producto?", "Debe revisar nombre, descripción, precio, imagen, estado y disponibilidad."),
            ("¿Qué pasa si detecta un problema de seguridad?", "Debe actuar de inmediato, restringir accesos y reportar la incidencia al equipo responsable."),
        ],
    },
    {
        "name": "Usuario sin registrar",
        "role_label": "usuario sin registrar",
        "role_description": "El usuario sin registrar puede ver el catálogo y conocer el sistema, pero debe crear una cuenta para realizar compras o acceder a funciones más completas.",
        "focus_points": [
            "Puede observar la plataforma y revisar el catálogo general.",
            "Puede aprender el funcionamiento del sistema antes de crear una cuenta.",
            "No puede completar compras ni usar funcionalidades completas hasta registrarse.",
        ],
        "faq_items": [
            ("¿Qué puedo hacer sin registrarme?", "Puede ver el catálogo y conocer la plataforma, pero no puede completar compras ni usar todas las funciones."),
            ("¿Por qué necesito registrar una cuenta?", "Porque el registro permite guardar datos, gestionar pedidos y acceder a funciones personalizadas."),
            ("¿Puedo volver después de ver la plataforma?", "Sí, puede registrarse en cualquier momento y continuar desde donde lo dejó."),
        ],
    },
]

for profile in profiles:
    doc = build_profile_manual(
        profile["name"],
        profile["role_label"],
        profile["role_description"],
        profile["focus_points"],
        profile["faq_items"],
    )
    filename = f"Manual_Usuario_YESA_{profile['name'].replace(' ', '_')}.docx"
    if profile["name"] == "Usuario sin registrar":
        filename = "Manual_Usuario_YESA_Usuario_Sin_Registrar.docx"
    output_path = out_dir / filename
    doc.save(str(output_path))
    print(f"Documento generado: {output_path}")

summary_doc = Document()
add_title(summary_doc, "Manual General de Usuario YESA", "Resumen completo para todos los perfiles", "Versión 2.0 • 2026")
add_paragraph(summary_doc, "Este documento resume las funciones principales de YESA para clientes, auxiliares, administradores y usuarios no registrados.")
add_bullets(summary_doc, [
    "Los clientes compran, revisan pedidos y usan favoritos.",
    "Los auxiliares apoyan tareas administrativas limitadas.",
    "Los administradores gestionan productos, usuarios, pedidos y soporte.",
    "Los usuarios sin registrar pueden explorar el sistema antes de crear una cuenta.",
])
summary_doc.save(str(out_dir / "Manual_Usuario_YESA_General.docx"))
print(f"Documento generado: {out_dir / 'Manual_Usuario_YESA_General.docx'}")
