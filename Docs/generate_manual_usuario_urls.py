from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = Path(__file__).resolve().parent.parent
out_dir = root / "Docs" / "documentacion final"
out_dir.mkdir(parents=True, exist_ok=True)


def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_title(doc, title, subtitle, version):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = 24

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = 12

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(version)
    run.italic = True
    run.font.size = 11
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    return table


def add_section(doc, title, paragraphs=None, bullets=None, level=1):
    doc.add_heading(title, level=level)
    if paragraphs:
        for paragraph in paragraphs:
            add_paragraph(doc, paragraph)
    if bullets:
        add_bullets(doc, bullets)


def get_page_url(title: str) -> str:
    mapping = {
        "Módulo de autenticación": "/login",
        "Pantalla de inicio de sesión": "/login",
        "Pantalla de registro": "/register",
        "Menú principal y navegación": "/",
        "Catálogo": "/",
        "Detalle de producto": "/producto/:id",
        "Carrito": "/carrito",
        "Checkout": "/checkout",
        "Personalización": "/personalizar",
        "Favoritos": "/favoritos",
        "Mis pedidos": "/mis-pedidos",
        "Mis cotizaciones": "/mis-cotizaciones",
        "Mis Consultas": "/mis-consultas",
        "Diseños guardados": "/disenos-guardados",
        "Perfil": "/perfil",
        "Dashboard administrativo": "/admin/dashboard",
        "Dashboard administrativo (auxiliar)": "/admin/dashboard",
        "Dashboard administrativo (administrador)": "/admin/dashboard",
        "Gestión de pedidos (auxiliar)": "/admin/pedidos",
        "Gestión de categorías": "/admin/categorias",
        "Gestión limitada de categorías": "/admin/categorias",
        "Gestión de subcategorías": "/admin/subcategorias",
        "Gestión limitada de subcategorias": "/admin/subcategorias",
        "Gestión de productos": "/admin/productos",
        "Gestión limitada de productos": "/admin/productos",
        "Soporte": "/admin/contactos",
        "Soporte y atención": "/admin/contactos",
    }
    return mapping.get(title, "")


def add_page_section(doc, title, description, bullets=None):
    doc.add_heading(title, level=2)
    add_paragraph(doc, description)
    url = get_page_url(title)
    if url:
        add_paragraph(doc, f"URL de la pantalla: {url}")
    if bullets:
        add_bullets(doc, bullets)


def get_role_screens(profile_name):
    common_pages = [
        {"title": "Módulo de autenticación", "description": "..."},
        {"title": "Pantalla de inicio de sesión", "description": "..."},
        {"title": "Pantalla de registro", "description": "..."},
    ]
    client_pages = [*common_pages, {"title": "Catálogo", "description": "..."}, {"title": "Detalle de producto", "description": "..."}]
    auxiliar_pages = [*client_pages, {"title": "Dashboard administrativo (auxiliar)", "description": "..."}, {"title": "Gestión de pedidos (auxiliar)", "description": "..."}]
    administrador_pages = [*auxiliar_pages, {"title": "Dashboard administrativo (administrador)", "description": "..."}, {"title": "Gestión de usuarios", "description": "..."}]
    if profile_name == "Usuario sin registrar":
        return [*common_pages]
    if profile_name == "Cliente":
        return client_pages
    if profile_name == "Auxiliar":
        return auxiliar_pages
    if profile_name == "Administrador":
        return administrador_pages
    return common_pages


def build_profile_manual(profile_name, role_label, role_description, focus_points, faq_items):
    doc = Document()
    add_title(
        doc,
        f"Manual de Usuario YESA - {profile_name}",
        f"Guía detallada para {role_label.lower()} del sistema YESA",
        "Versión 2.0 • 2026 • Elaborado para usuarios sin conocimientos técnicos",
    )
    add_paragraph(doc, "ÍNDICE")
    add_bullets(doc, ["1. Descripción del sistema", "6. Pantallas clave y botones"])
    add_section(doc, "1. Descripción del sistema", [f"Este documento está orientado a {role_label}."])
    doc.add_heading("6. Pantallas clave y botones", level=1)
    for page in get_role_screens(profile_name):
        add_page_section(doc, page["title"], page.get("description", ""), page.get("bullets"))
    return doc


profiles = [
    {"name": "Cliente", "role_label": "cliente", "role_description": "...", "focus_points": [], "faq_items": []},
    {"name": "Auxiliar", "role_label": "auxiliar", "role_description": "...", "focus_points": [], "faq_items": []},
    {"name": "Administrador", "role_label": "administrador", "role_description": "...", "focus_points": [], "faq_items": []},
    {"name": "Usuario sin registrar", "role_label": "usuario sin registrar", "role_description": "...", "focus_points": [], "faq_items": []},
]

for profile in profiles:
    doc = build_profile_manual(
        profile["name"],
        profile["role_label"],
        profile["role_description"],
        profile["focus_points"],
        profile["faq_items"],
    )
    filename = f"Manual_Usuario_YESA_{profile['name'].replace(' ', '_')}.docx"
    if profile["name"] == "Usuario sin registrar":
        filename = "Manual_Usuario_YESA_Usuario_Sin_Registrar.docx"
    output_path = out_dir / filename
    doc.save(str(output_path))
    print(f"Documento generado: {output_path}")
