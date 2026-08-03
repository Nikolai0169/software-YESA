from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import datetime

root = Path(__file__).resolve().parent.parent
template_path = root / "Docs" / "plantillas" / "Plan Capacitacion Plantilla.docx"
out_dir = root / "Docs" / "documentacion final"
out_dir.mkdir(parents=True, exist_ok=True)


def clear_document(doc):
    body = doc.element.body
    for element in list(body):
        body.remove(element)


def add_title(doc, title, subtitle, version):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = doc.styles['Normal'].font.size

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = doc.styles['Normal'].font.size

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(version)
    run.italic = True
    run.font.size = doc.styles['Normal'].font.size
    doc.add_paragraph()


def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.add_run(f"• {item}")


def add_numbered(doc, items):
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{index}. {item}")


def set_paragraph_text(paragraph, text):
    paragraph.text = text


def find_paragraph(doc, contains):
    return next((p for p in doc.paragraphs if contains in p.text), None)


def insert_paragraph_after(paragraph, text, style='Normal'):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_run = new_para.add_run(text)
    new_para.style = style
    return new_para


def insert_paragraphs_after(paragraph, texts, style='Normal'):
    current = paragraph
    for text in texts:
        current = insert_paragraph_after(current, text, style=style)
    return current


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def get_paragraph_index(doc, paragraph):
    for index, p in enumerate(doc.paragraphs):
        if p._p is paragraph._p:
            return index
    return None


def remove_paragraphs_after(doc, paragraph):
    start = get_paragraph_index(doc, paragraph)
    if start is None:
        return
    start += 1
    for p in reversed(doc.paragraphs[start:]):
        remove_paragraph(p)


def fill_control_table(table, project_name, document_name, version, creation_date, responsible, leader):
    def set_cells(row, value):
        for col_index in range(2, len(row.cells)):
            row.cells[col_index].text = value

    set_cells(table.rows[1], project_name)
    set_cells(table.rows[2], document_name)
    set_cells(table.rows[3], version)
    set_cells(table.rows[4], creation_date)
    set_cells(table.rows[5], creation_date)
    set_cells(table.rows[6], responsible)
    set_cells(table.rows[7], leader)

    # Clear the history placeholder row and any trailing empty row in the control table
    for cell in table.rows[10].cells:
        cell.text = ''
    for cell in table.rows[11].cells:
        cell.text = ''


def fill_role_table(table, role_name):
    if len(table.rows) > 1:
        table.rows[1].cells[0].text = 'Rol'
        table.rows[1].cells[1].text = role_name
        for row in table.rows[2:]:
            row.cells[0].text = ''
            row.cells[1].text = ''


def fill_budget_table(table):
    budget_fills = {
        2: ['Lápices', 'Unid.', '20', '1.5', '30'],
        3: ['Esferos', 'Caja', '3', '20', '60'],
        4: ['Papel A4', 'Resma', '1', '20', '20'],
        5: ['Carpetas', 'Docena', '2', '15', '30'],
        7: ['Documentos', '-', '1', '20', '20'],
        9: ['Servicios profesionales', 'Aseo', '1', '200', '200'],
        11: ['Proyector Multimedia', 'Alquiler', '2', '60', '120'],
    }

    for row_index, row_data in budget_fills.items():
        if row_index < len(table.rows):
            for col_index, value in enumerate(row_data):
                table.rows[row_index].cells[col_index].text = value

    # Keep category headers and totals as defined by template, but ensure totals are updated.
    if len(table.rows) > 14:
        table.rows[14].cells[4].text = '430'
    if len(table.rows) > 15:
        table.rows[15].cells[4].text = '43'
    if len(table.rows) > 16:
        table.rows[16].cells[4].text = '473'


def build_training_plan(role_name, role_focus):
    doc = Document(template_path)
    project_name = 'YESA - Plataforma de comercio electrónico'
    document_name = f'Plan de Capacitación YESA - {role_name}'
    version = '1.0'
    creation_date = datetime.date.today().isoformat()
    responsible = 'Equipo de implementación YESA'
    leader = 'Coordinador de capacitación'

    fill_control_table(doc.tables[0], project_name, document_name, version, creation_date, responsible, leader)
    fill_role_table(doc.tables[2], role_name)
    fill_budget_table(doc.tables[3])

    replace_paragraph_text = find_paragraph
    toc_para = find_paragraph(doc, '<<Elabore la tabla de contenido del plan')
    if toc_para:
        set_paragraph_text(toc_para, '1. Introducción')
        insert_paragraphs_after(toc_para, [
            '2. Presentación del sistema',
            '3. Convenciones',
            '4. Actividades y funciones de la capacitación',
            '5. Propuesta de agenda por perfil',
            '6. Revisión logística',
            '7. Recursos y presupuesto',
            '8. Evaluación y seguimiento',
            '9. Anexos',
        ], style='Normal')

    intro_para = find_paragraph(doc, 'El Plan de capacitación tiene como objetivo')
    if intro_para:
        set_paragraph_text(intro_para, f'Este Plan de Capacitación está diseñado para guiar el uso de YESA por el rol de {role_name}. Su propósito es asegurar que cada usuario comprenda sus responsabilidades, conozca los flujos operativos y aplique buenas prácticas de seguridad y atención al cliente.')

    follow_intro_para = find_paragraph(doc, 'La capacitación uniformiza criterios sensibilizando')
    if follow_intro_para:
        set_paragraph_text(follow_intro_para, 'La capacitación busca uniformizar criterios y generar confianza en el uso de la plataforma, explicando los límites de cada rol y la importancia de proteger datos sensibles y procesos administrativos.')

    definitions_para = find_paragraph(doc, '<<Definición de todos los términos')
    if definitions_para:
        set_paragraph_text(definitions_para, 'Catálogo: Conjunto de productos visibles para los usuarios en la plataforma.\nCarrito: Espacio donde se agregan los productos antes de confirmar la compra.\nPedido: Compra confirmada registrada en el sistema con estado y seguimiento.\nCotización: Solicitud de precio para un diseño personalizado o un producto especial.\nDiseño guardado: Diseño 3D creado por el cliente y almacenado para edición o cotización posterior.\nRol: Perfil de usuario que define permisos y accesos en el sistema.\nSoporte: Canal de atención para consultas, problemas o solicitudes de ayuda.')

    presentation_para = find_paragraph(doc, '<<Escriba la presentación del software')
    if presentation_para:
        set_paragraph_text(presentation_para, 'YESA es una plataforma de comercio electrónico que integra catálogo de productos, carrito de compras, personalización 3D, cotizaciones y gestión administrativa. Está diseñada para permitir compras seguras, gestión de pedidos y atención por roles diferenciados. El sistema cuenta con tres perfiles principales: Cliente, Auxiliar y Administrador. Cada perfil tiene acceso a funciones específicas según su responsabilidad y autorización.')

    conventions_para = find_paragraph(doc, '<<Si utiliza dentro del manual algunas convenciones')
    if conventions_para:
        set_paragraph_text(conventions_para, 'En este documento se utilizan las siguientes convenciones:\n• Nombres de botones entre comillas simples, por ejemplo: botón "Iniciar Sesión".\n• Rutas y menús descritos con su nombre exacto según la interfaz del sistema.\n• Tareas prácticas asignadas según el rol correspondiente.')

    if conventions_para:
        remove_paragraphs_after(doc, conventions_para)
        current = conventions_para
        current = insert_paragraphs_after(current, [
            'ACTIVIDADES Y FUNCIONES DE LA CAPACITACIÓN',
            'La capacitación se divide en tres fases principales: preparación, demostración y práctica por rol. Cada etapa incluye objetivos claros, tiempos estimados y resultados esperados.',
            'ETAPA 1: PREPARACIÓN Y CONTEXTO',
            'Objetivo de la actividad: Presentar la plataforma YESA, explicar los roles y establecer el entorno de trabajo para la capacitación.',
        ], style='Título 1')
        current = insert_paragraphs_after(current, [
            '• Revisión del propósito de YESA y el alcance funcional del sistema.',
            '• Identificación de los roles: Administrador, Auxiliar y Cliente.',
            '• Explicación de la importancia de la seguridad y permisos.',
            '• Verificación de accesos y datos de prueba.',
        ], style='Normal')
        current = insert_paragraphs_after(current, [
            'ETAPA 2: DEMOSTRACIÓN GUIADA DEL SISTEMA',
            'Objetivo de la actividad: Mostrar los flujos principales para cada rol y facilitar la comprensión de las pantallas clave.',
        ], style='Título 1')
        current = insert_paragraphs_after(current, [
            '• Demostración del proceso de registro e inicio de sesión.',
            '• Mostrar catálogo, carrito y checkout para el Cliente.',
            '• Presentación del dashboard administrativo y gestión de pedidos para Auxiliar.',
            '• Revisión de usuarios, cotizaciones y soporte para Administrador.',
        ], style='Normal')
        current = insert_paragraphs_after(current, [
            'ETAPA 3: PRÁCTICA POR ROL',
            'Objetivo de la actividad: Ejecutar ejercicios prácticos específicos para cada rol.',
        ], style='Título 1')
        role_exercises = {
            'Cliente': [
                'Registro en la plataforma y acceso seguro.',
                'Búsqueda de productos, uso del carrito y finalización de la compra.',
                'Revisión de pedidos y uso de soporte.',
            ],
            'Auxiliar': [
                'Acceso al dashboard administrativo y revisión de pedidos.',
                'Consulta y actualización de estados de pedidos.',
                'Revisión de productos y categorías con acceso limitado.',
            ],
            'Administrador': [
                'Gestión de usuarios y revisión de permisos.',
                'Creación y edición de categorías, subcategorías y productos.',
                'Revisión de cotizaciones y gestión de soporte.',
            ],
        }
        exercises = role_exercises.get(role_name, [
            'Acceder al sistema y revisar las funciones asignadas al rol.',
        ])
        current = insert_paragraphs_after(current, [f'• {item}' for item in exercises], style='Normal')
        current = insert_paragraphs_after(current, [
            'PROPUESTA DE AGENDA POR PERFIL',
            'La capacitación se organiza en sesiones específicas por perfil, con tiempos orientativos y actividades prácticas necesarias para cada grupo.',
        ], style='Título 1')
        current = insert_paragraphs_after(current, [
            'Sesión 1: Presentación del sistema y acceso al entorno de pruebas (30 min).',
            f'Sesión 2: Demostración de los flujos clave para el rol de {role_name} (60 min).',
            'Sesión 3: Ejercicio práctico guiado por rol (60 min).',
            'Sesión 4: Evaluación, dudas y cierre (30 min).',
        ], style='Normal')
        current = insert_paragraphs_after(current, [
            'REVISIÓN LOGÍSTICA',
            'Esta sección describe los elementos necesarios para llevar a cabo la capacitación de manera efectiva.',
        ], style='Título 1')
        current = insert_paragraphs_after(current, [
            '• Sala con proyector o pantalla para demostraciones en vivo.',
            '• Computadoras con navegador actualizado para cada participante o estación compartida.',
            '• Conexión a internet estable y acceso al entorno de pruebas de YESA.',
            '• Credenciales de usuario de prueba para Cliente, Auxiliar y Administrador.',
            '• Manual de usuario impreso o digital para referencia durante la sesión.',
        ], style='Normal')
        current = insert_paragraphs_after(current, [
            'RECURSOS Y PRESUPUESTO',
            'Los recursos incluyen herramientas, materiales y personal requerido para ejecutar la capacitación.',
        ], style='Título 1')
        current = insert_paragraphs_after(current, [
            '• Facilitador de capacitación con conocimiento de la plataforma YESA.',
            '• Equipo de soporte técnico para resolver problemas de acceso o conectividad.',
            '• Documentación de soporte impresa o en formato digital.',
            '• Tiempo del equipo: estimado en 9 horas de capacitación total para los tres perfiles.',
        ], style='Normal')
        current = insert_paragraphs_after(current, [
            'EVALUACIÓN Y SEGUIMIENTO',
            'La evaluación se realiza mediante ejercicios prácticos y observación del uso correcto del sistema.',
        ], style='Título 1')
        current = insert_paragraphs_after(current, [
            '• Lista de verificación de tareas completadas por rol.',
            '• Evaluación de resultados práctica: al menos 3 tareas clave ejecutadas correctamente.',
            '• Revisión de dudas y errores comunes al final de cada sesión.',
            '• Sesión de seguimiento 7 días después para responder preguntas y reforzar el uso correcto.',
        ], style='Normal')
        current = insert_paragraphs_after(current, [
            'ANEXOS',
            'Los anexos pueden incluir formularios de evaluación, capturas de pantalla del sistema, ejemplos de flujos y contactos de soporte.',
        ], style='Título 1')
        current = insert_paragraphs_after(current, [
            '• Formulario de evaluación de la capacitación.',
            '• Lista de verificación de pasos clave por rol.',
            '• Contacto del responsable de soporte técnico.',
        ], style='Normal')

    return doc


def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = doc.styles['Normal'].font.size
    return p


if __name__ == '__main__':
    roles = [
        ('Cliente', 'Usuario final que compra, revisa pedidos y usa soporte.'),
        ('Auxiliar', 'Rol de apoyo operativo con acceso restringido a funciones administrativas.'),
        ('Administrador', 'Rol de gestión completa que administra productos, usuarios, pedidos y soporte.'),
    ]
    for role_name, role_focus in roles:
        plan_doc = build_training_plan(role_name, role_focus)
        output_path = out_dir / f'Plan_Capacitacion_YESA_{role_name}.docx'
        plan_doc.save(str(output_path))
        print(f'Documento generado: {output_path}')
